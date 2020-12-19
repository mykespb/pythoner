# работа с  офисными документами,
# создание с начала и с красивым оформлением
# myke. 2020-12-19 0.4

import docx
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

# -------------------------------------------------------- make1

def make1():
    """ делаем простой документ, но уже с разметкой.
    это helloworld.docx 
    """

    doc = docx.Document()

    # добавляем первый параграф
    doc.add_paragraph('Здравствуй, мир!')

    # добавляем еще два параграфа
    par1 = doc.add_paragraph('Это второй абзац.')
    par2 = doc.add_paragraph('Это третий абзац.')

    # добавляем текст во второй параграф
    par1.add_run(' Этот текст был добавлен во второй абзац.')

    # добавляем текст в третий параграф
    par2.add_run(' Добавляем текст в третий абзац.').bold = True

    # делаем заголовки
    doc.add_heading('Заголовок 0', 0)
    doc.add_heading('Заголовок 1', 1)
    doc.add_heading('Заголовок 2', 2)
    doc.add_heading('Заголовок 3', 3)
    doc.add_heading('Заголовок 4', 4)

    # запись результата
    doc.save('helloworld.docx')

# -------------------------------------------------------- make1

def make2():
    """ делаем документ с таблицей
    это table.docx 
    """

    doc = docx.Document()

    # добавляем таблицу 3x3
    table = doc.add_table(rows = 3, cols = 3)
    # применяем стиль для таблицы
    table.style = 'Table Grid'

    # заполняем таблицу данными
    for row in range(3):
        for col in range(3):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            cell.text = str(row + 1) + str(col + 1)

    doc.save('table.docx')

    doc = docx.Document('table.docx')

    # получаем первую таблицу в документе
    table = doc.tables[0]

    # читаем данные из таблицы
    for row in table.rows:
        string = ''
        for cell in row.cells:
            string = string + cell.text + ' '
        print(string)

# -------------------------------------------------------- make1

def make3():
    """ делаем документ с таблицей
    это demo.docx 
    """

    photo = 'photo1.jpg'

    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture(photo, width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)
    merged = c1.merge(c2)
    merged.text = "NEW!"

    document.add_page_break()

    document.save('demo.docx')

#make1()
#make2()
make3()
