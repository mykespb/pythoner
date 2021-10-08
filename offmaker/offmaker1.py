#!python3

# (C) Mikhail Kolodin, 2021
# offmaker.py 2021-10-03 2021-10-09 1.3

# обработка вордовых шаблонов и эксельных списков
# с получением вордовых писем в отдельных файлах

import os
import openpyxl
import docx

template  = "template.docx"
lopersons = "lopersons.xlsx"
outdir    = "outdir"

print ("here we start with {} as template and {} as list of persons\nresult files will be written to {} directory".format
    (template, lopersons, outdir))

try:
    os.mkdir (outdir)
except:
    pass

wb = openpyxl.load_workbook (lopersons)
sheet = wb.active
print (sheet)

print("Rows: {}, columns: {}" . format(smr := sheet.max_row, smc := sheet.max_column))
cols = []
for item in sheet.iter_cols(0, smc):
    cols.append(item[0].value)
print(cols)

data = {}

doc = docx.Document(template)

print("было:")
for para in doc.paragraphs:
    txt = para.text
    print (txt)

for i, row in enumerate (sheet.iter_rows (values_only = True)):
    if i == 0:
        continue

    print ("-------------------------------\nperson {}" .format(i))
    for j in range(smc):
        print ("data {} is {}" .format(cols[j], row[j]))
        data[cols[j]] = row[j]
    print(data)

    outname = f"./{outdir}/{i:03d}.docx"
    print (f"will write to file: {outname}")

    outdoc = docx.Document()

    print ("стало:")
    for para in doc.paragraphs:
        txt = para.text
        for de in data:
            txt = txt.replace ("{"+de+"}", str(data[de]))
        print (txt)
        doc_para = outdoc.add_paragraph(txt)

    outdoc.save(outname)

# useful links:

# https://www.marsja.se/your-guide-to-reading-excel-xlsx-files-in-python/
# Your Guide to Reading Excel (xlsx) Files in Python
# by Erik Marsja | Feb 16, 2020 | Programming, Python
